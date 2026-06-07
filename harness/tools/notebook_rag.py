from __future__ import annotations

import hashlib
import inspect
import math
import re
import struct
from pathlib import Path
from typing import Any

from harness import config, db


_MODEL: Any = None
_MODEL_NAME: str | None = None


IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tif", ".tiff"}
AUDIO_EXTS = {".wav", ".mp3", ".m4a", ".flac", ".ogg", ".aac", ".wma"}


def _embed_dim() -> int:
    try:
        return int(getattr(config, "EMBED_DIM", 384))
    except (TypeError, ValueError):
        return 384


def _kind_from_path(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix in {".xlsx", ".xlsm"}:
        return "xlsx"
    if suffix == ".md":
        return "md"
    if suffix == ".txt":
        return "txt"
    if suffix in IMAGE_EXTS:
        return "image"
    if suffix in AUDIO_EXTS:
        return "audio"
    return suffix.lstrip(".") or "file"


def _call_by_signature(fn: Any, **values: Any) -> Any:
    try:
        signature = inspect.signature(fn)
    except (TypeError, ValueError):
        return fn(**values)

    parameters = signature.parameters
    if any(param.kind == inspect.Parameter.VAR_KEYWORD for param in parameters.values()):
        return fn(**values)

    kwargs = {
        name: values[name]
        for name, param in parameters.items()
        if name in values
        and param.kind
        in {
            inspect.Parameter.POSITIONAL_OR_KEYWORD,
            inspect.Parameter.KEYWORD_ONLY,
        }
    }
    if kwargs:
        return fn(**kwargs)
    return fn()


def _add_source(notebook_id: int, path: Path, kind: str) -> Any:
    values = {
        "notebook_id": notebook_id,
        "path": str(path),
        "kind": kind,
        "status": "ingesting",
        "title": path.name,
        "name": path.name,
    }
    try:
        return _call_by_signature(db.add_source, **values)
    except TypeError:
        for args in (
            (notebook_id, str(path), kind),
            (notebook_id, str(path), kind, "ingesting"),
            (notebook_id, str(path)),
        ):
            try:
                return db.add_source(*args)
            except TypeError:
                continue
        raise


def _source_id(source: Any) -> Any:
    if isinstance(source, dict):
        return source.get("source_id") or source.get("id")
    for attr in ("source_id", "id"):
        if hasattr(source, attr):
            return getattr(source, attr)
    return source


def _set_source_status(source_id: Any, status: str, error: str | None = None) -> None:
    values = {"source_id": source_id, "id": source_id, "status": status, "error": error}
    try:
        _call_by_signature(db.set_source_status, **values)
    except TypeError:
        for args in ((source_id, status, error), (source_id, status)):
            try:
                db.set_source_status(*args)
                return
            except TypeError:
                continue
        raise


def _add_chunks(notebook_id: int, source_id: Any, chunks: list[dict[str, Any]]) -> None:
    values = {"notebook_id": notebook_id, "source_id": source_id, "id": source_id, "chunks": chunks}
    try:
        _call_by_signature(db.add_chunks, **values)
    except TypeError:
        for args in ((source_id, chunks), (notebook_id, source_id, chunks), (notebook_id, chunks)):
            try:
                db.add_chunks(*args)
                return
            except TypeError:
                continue
        raise


def _get_chunks(notebook_id: int) -> list[Any]:
    values = {"notebook_id": notebook_id}
    try:
        chunks = _call_by_signature(db.get_chunks, **values)
    except TypeError:
        chunks = db.get_chunks(notebook_id)
    return list(chunks or [])


def _field(item: Any, name: str, default: Any = None) -> Any:
    if isinstance(item, dict):
        return item.get(name, default)
    try:
        return item[name]
    except Exception:
        return getattr(item, name, default)


def _normalize(vec: list[float]) -> list[float]:
    norm = math.sqrt(sum(value * value for value in vec))
    if norm <= 0:
        return vec
    return [value / norm for value in vec]


def _hash_embedding(text: str, dim: int) -> list[float]:
    vec = [0.0] * dim
    tokens = re.findall(r"[\w가-힣]+", text.lower())
    if not tokens:
        tokens = [text[:64] or "empty"]
    for token in tokens:
        digest = hashlib.blake2b(token.encode("utf-8", "ignore"), digest_size=8).digest()
        number = int.from_bytes(digest, "little", signed=False)
        index = number % dim
        sign = -1.0 if (number >> 63) else 1.0
        vec[index] += sign
    return _normalize(vec)


def _embed(texts: list[str]) -> list[list[float]]:
    dim = _embed_dim()
    model_name = str(getattr(config, "EMBED_MODEL", "") or "")

    global _MODEL, _MODEL_NAME
    try:
        from sentence_transformers import SentenceTransformer

        if _MODEL is None or _MODEL_NAME != model_name:
            _MODEL = SentenceTransformer(model_name)
            _MODEL_NAME = model_name
        encoded = _MODEL.encode(texts, convert_to_numpy=False, show_progress_bar=False)
        vectors: list[list[float]] = []
        for vector in encoded:
            values = [float(value) for value in vector]
            vectors.append(_normalize(values))
        return vectors
    except Exception:
        return [_hash_embedding(text, dim) for text in texts]


def _embedding_to_bytes(vector: list[float]) -> bytes:
    try:
        import numpy as np

        return np.asarray(vector, dtype=np.float32).tobytes()
    except ImportError:
        return struct.pack("<" + "f" * len(vector), *[float(value) for value in vector])


def _embedding_from_blob(value: Any) -> list[float]:
    if value is None:
        return []
    if isinstance(value, list):
        return [float(item) for item in value]
    if isinstance(value, tuple):
        return [float(item) for item in value]
    if isinstance(value, memoryview):
        value = value.tobytes()
    if isinstance(value, bytearray):
        value = bytes(value)
    if isinstance(value, bytes):
        if len(value) % 4:
            return []
        try:
            import numpy as np

            return [float(item) for item in np.frombuffer(value, dtype=np.float32)]
        except ImportError:
            count = len(value) // 4
            return [float(item) for item in struct.unpack("<" + "f" * count, value)]
    return []


def _cosine(left: list[float], right: list[float]) -> float:
    if not left or not right:
        return 0.0
    size = min(len(left), len(right))
    dot = sum(left[index] * right[index] for index in range(size))
    left_norm = math.sqrt(sum(value * value for value in left[:size]))
    right_norm = math.sqrt(sum(value * value for value in right[:size]))
    if left_norm <= 0 or right_norm <= 0:
        return 0.0
    return dot / (left_norm * right_norm)


def _read_text(path: Path) -> list[tuple[str, str]]:
    return [("file", path.read_text(encoding="utf-8", errors="ignore"))]


def _read_pdf(path: Path) -> list[tuple[str, str]] | str:
    try:
        import fitz
    except ImportError:
        return "error: 이 기능은 pymupdf 설치가 필요합니다."

    parts: list[tuple[str, str]] = []
    with fitz.open(str(path)) as document:
        for page_index, page in enumerate(document, start=1):
            text = page.get_text("text").strip()
            if text:
                parts.append((f"p.{page_index}", text))
    return parts


def _read_docx(path: Path) -> list[tuple[str, str]] | str:
    try:
        from docx import Document
    except ImportError:
        return "error: 이 기능은 python-docx 설치가 필요합니다."

    document = Document(str(path))
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return [("docx", "\n\n".join(parts))]


def _read_xlsx(path: Path) -> list[tuple[str, str]] | str:
    try:
        import openpyxl
    except ImportError:
        return "error: 이 기능은 openpyxl 설치가 필요합니다."

    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts: list[tuple[str, str]] = []
    try:
        for sheet in workbook.worksheets:
            rows: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                values = [str(value) for value in row if value is not None and str(value).strip()]
                if values:
                    rows.append(" | ".join(values))
            if rows:
                parts.append((sheet.title, "\n".join(rows)))
    finally:
        workbook.close()
    return parts


def _extract(path: Path, kind: str) -> list[tuple[str, str]] | str:
    if kind == "pdf":
        return _read_pdf(path)
    if kind == "docx":
        return _read_docx(path)
    if kind == "xlsx":
        return _read_xlsx(path)
    if kind in {"txt", "md"}:
        return _read_text(path)
    if kind in {"image", "audio"}:
        return [(kind, f"{path.name}: Gemma-native {kind} processing is pending.")]
    return _read_text(path)


def _split_long(text: str, target: int, overlap: int) -> list[str]:
    chunks: list[str] = []
    step = max(1, target - overlap)
    for start in range(0, len(text), step):
        chunk = text[start : start + target].strip()
        if chunk:
            chunks.append(chunk)
        if start + target >= len(text):
            break
    return chunks


def _chunk_segment(locator: str, text: str, target: int = 800, overlap: int = 100) -> list[tuple[str, str]]:
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
    chunks: list[tuple[str, str]] = []
    current = ""

    for paragraph in paragraphs:
        if len(paragraph) > target:
            if current.strip():
                chunks.append((locator, current.strip()))
                current = ""
            for piece in _split_long(paragraph, target, overlap):
                chunks.append((locator, piece))
            continue

        next_text = paragraph if not current else current + "\n\n" + paragraph
        if len(next_text) > target and current:
            chunks.append((locator, current.strip()))
            tail = current[-overlap:].strip()
            current = (tail + "\n\n" + paragraph).strip() if tail else paragraph
        else:
            current = next_text

    if current.strip():
        chunks.append((locator, current.strip()))
    return chunks


def _chunk_extracted(parts: list[tuple[str, str]]) -> list[tuple[str, str]]:
    chunks: list[tuple[str, str]] = []
    for locator, text in parts:
        chunks.extend(_chunk_segment(locator, text))
    return chunks


def ingest_source(notebook_id: int, path: str) -> dict[str, Any]:
    source_path = Path(path)
    if not source_path.exists():
        return {"ok": False, "error": f"파일을 찾을 수 없습니다: {path}"}

    kind = _kind_from_path(source_path)
    source = _add_source(notebook_id, source_path, kind)
    source_id = _source_id(source)

    try:
        _set_source_status(source_id, "ingesting")
        extracted = _extract(source_path, kind)
        if isinstance(extracted, str):
            _set_source_status(source_id, "error", extracted)
            return {"ok": False, "error": extracted, "source_id": source_id}

        text_chunks = _chunk_extracted(extracted)
        if not text_chunks:
            _set_source_status(source_id, "error", "추출된 텍스트가 없습니다.")
            return {"ok": False, "error": "추출된 텍스트가 없습니다.", "source_id": source_id}

        vectors = _embed([text for _, text in text_chunks])
        chunks = [
            {
                "ordinal": index,
                "text": text,
                "locator": locator,
                "embedding": _embedding_to_bytes(vectors[index]),
            }
            for index, (locator, text) in enumerate(text_chunks)
        ]
        _add_chunks(notebook_id, source_id, chunks)
        _set_source_status(source_id, "ready")
        return {"ok": True, "source_id": source_id, "n_chunks": len(chunks)}
    except Exception as exc:
        try:
            _set_source_status(source_id, "error", str(exc))
        except Exception:
            pass
        return {"ok": False, "error": str(exc), "source_id": source_id}


def context_for(notebook_id: int, query: str, k: int = 5) -> str:
    try:
        limit = max(1, min(20, int(k)))
    except (TypeError, ValueError):
        limit = 5

    chunks = _get_chunks(notebook_id)
    if not chunks:
        return ""

    query_vector = _embed([query])[0]
    scored: list[tuple[float, Any]] = []
    for chunk in chunks:
        vector = _embedding_from_blob(_field(chunk, "embedding"))
        score = _cosine(query_vector, vector)
        if score > 0:
            scored.append((score, chunk))

    if not scored:
        return ""

    scored.sort(key=lambda item: item[0], reverse=True)
    blocks = ["[소스 발췌]"]
    for index, (_, chunk) in enumerate(scored[:limit], start=1):
        locator = _field(chunk, "locator", "unknown")
        text = str(_field(chunk, "text", "")).strip()
        if text:
            blocks.append(f"[{index}] {locator}\n{text}")
    return "\n\n".join(blocks) if len(blocks) > 1 else ""


def notebook_search(notebook_id: int, query: str, k: int = 5) -> str:
    return context_for(notebook_id, query, k)


def register(registry: Any) -> None:
    registry.add(
        "notebook_search",
        "노트북에 저장된 자료에서 관련 발췌를 검색합니다.",
        {
            "type": "object",
            "properties": {
                "notebook_id": {"type": "integer"},
                "query": {"type": "string"},
                "k": {"type": "integer", "default": 5, "minimum": 1, "maximum": 20},
            },
            "required": ["notebook_id", "query"],
        },
        notebook_search,
        permissions=(),
    )

